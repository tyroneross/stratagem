"""Word document tools — read .docx files."""

from typing import Any
from pathlib import Path

from claude_agent_sdk import tool


@tool(
    "read_docx",
    "Read a Word document (.docx) and extract text, tables, and metadata as markdown.",
    {
        "type": "object",
        "properties": {
            "file_path": {"type": "string", "description": "Absolute path to the .docx file"},
            "include_tables": {"type": "boolean", "description": "Extract tables as markdown", "default": True},
            "include_metadata": {"type": "boolean", "description": "Include document metadata", "default": True},
        },
        "required": ["file_path"],
    },
)
async def read_docx(args: dict[str, Any]) -> dict[str, Any]:
    file_path = args["file_path"]
    include_tables = args.get("include_tables", True)
    include_metadata = args.get("include_metadata", True)

    try:
        from docx import Document
    except ImportError:
        return _error("python-docx not installed. Run: pip install python-docx")

    path = Path(file_path)
    if not path.exists():
        return _error(f"File not found: {file_path}")

    try:
        doc = Document(file_path)
    except Exception as e:
        return _error(f"Failed to open DOCX: {e}")

    sections: list[str] = [
        f"# Document: {path.name}",
        "",
    ]

    # Metadata
    if include_metadata:
        props = doc.core_properties
        meta_parts = []
        if props.title:
            meta_parts.append(f"**Title**: {props.title}")
        if props.author:
            meta_parts.append(f"**Author**: {props.author}")
        if props.created:
            meta_parts.append(f"**Created**: {props.created.strftime('%Y-%m-%d')}")
        if props.modified:
            meta_parts.append(f"**Modified**: {props.modified.strftime('%Y-%m-%d')}")
        if meta_parts:
            sections.extend(meta_parts)
            sections.append("")

    # Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""

        if "Heading 1" in style_name:
            sections.append(f"## {text}")
        elif "Heading 2" in style_name:
            sections.append(f"### {text}")
        elif "Heading 3" in style_name:
            sections.append(f"#### {text}")
        elif "List Bullet" in style_name:
            sections.append(f"- {text}")
        elif "List Number" in style_name:
            sections.append(f"1. {text}")
        else:
            sections.append(text)

        sections.append("")

    # Tables
    if include_tables and doc.tables:
        sections.append("---")
        sections.append(f"## Tables ({len(doc.tables)} found)")
        sections.append("")

        for t_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                rows.append(cells)

            if rows:
                sections.append(f"### Table {t_idx + 1}")
                sections.append(_rows_to_markdown(rows))
                sections.append("")

    return {"content": [{"type": "text", "text": "\n".join(sections)}]}


def _rows_to_markdown(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    col_count = max(len(r) for r in rows)
    lines = []
    for i, row in enumerate(rows):
        padded = row + [""] * (col_count - len(row))
        cleaned = [cell.replace("|", "\\|") for cell in padded]
        lines.append("| " + " | ".join(cleaned) + " |")
        if i == 0:
            lines.append("| " + " | ".join(["---"] * col_count) + " |")
    return "\n".join(lines)


def _error(msg: str) -> dict[str, Any]:
    return {"content": [{"type": "text", "text": f"Error: {msg}"}], "isError": True}
